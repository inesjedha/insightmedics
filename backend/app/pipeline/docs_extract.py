"""Extraction texte des documents méthodologiques (protocole, questionnaire)."""

from __future__ import annotations

from pathlib import Path


def extract_text(path: str | Path) -> str:
    path = Path(path)
    ext = path.suffix.lower()
    if ext == ".docx":
        import docx

        d = docx.Document(str(path))
        parts = [p.text for p in d.paragraphs if p.text.strip()]
        for table in d.tables:
            for row in table.rows:
                parts.append(" | ".join(c.text.strip() for c in row.cells))
        return "\n".join(parts)
    if ext == ".pdf":
        import fitz  # pymupdf

        with fitz.open(str(path)) as doc:
            return "\n".join(page.get_text() for page in doc)
    if ext in (".txt", ".md"):
        return path.read_text(encoding="utf-8", errors="replace")
    raise ValueError(f"Format de document non supporté : {ext} (.docx, .pdf, .txt)")
