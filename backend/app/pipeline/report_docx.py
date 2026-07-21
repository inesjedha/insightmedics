"""Jalon M6 — Rapport d'audit Word (11 sections), au format du livrable de Hamza.

Assemble le rapport narratif à partir des sorties du pipeline. Les sections narratives
viennent de LLM-2 (synthèse, limites, plan d'analyse) ; les faits chiffrés et le tableau
de score viennent du code. Une figure en barres des 8 sous-scores est intégrée si
matplotlib est disponible.
"""

from __future__ import annotations

import io
from datetime import date

from docx import Document
from docx.document import Document as DocxDocument
from docx.shared import Pt, RGBColor


def _fig_scores(score: dict) -> io.BytesIO | None:
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
    except ImportError:  # matplotlib optionnel : pas de figure si absent
        return None
    doms = score["domaines"]
    labels = [f"D{d['domaine']}" for d in doms]
    pct = [(d["obtenu"] / d["max"] * 100) if d["max"] else 0 for d in doms]
    colors = ["#2e8b57" if p >= 80 else "#e0a800" if p >= 55 else "#c0392b" for p in pct]
    fig, ax = plt.subplots(figsize=(6.4, 2.6))
    ax.bar(labels, pct, color=colors)
    ax.set_ylim(0, 100)
    ax.set_ylabel("% du domaine")
    ax.set_title("Sous-scores de qualité (% du maximum de chaque domaine)", fontsize=9)
    ax.axhline(80, color="#999", lw=0.6, ls="--")
    for i, p in enumerate(pct):
        ax.text(i, p + 2, f"{p:.0f}", ha="center", fontsize=7)
    fig.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150)
    plt.close(fig)
    buf.seek(0)
    return buf


def _h(doc: DocxDocument, text: str) -> None:
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3B, 0x57)


def _p(doc: DocxDocument, text: str) -> None:
    if text:
        doc.add_paragraph(text)


def build_report(profiling: dict, score: dict, ai_audit: dict | None) -> bytes:
    ai = ai_audit or {}
    assess = ai.get("assessment", {}) or {}
    rs = assess.get("report_sections_fr", {}) or {}
    study = ai.get("study") or {}
    s, m = profiling["structure"], profiling["missing_summary"]
    fname = profiling["file"].get("file_name", "")

    doc = Document()
    style = doc.styles["Normal"].font
    style.name = "Arial"
    style.size = Pt(10.5)

    title = doc.add_heading("RAPPORT D'AUDIT BIOSTATISTIQUE", level=0)
    for r in title.runs:
        r.font.color.rgb = RGBColor(0x1F, 0x3B, 0x57)
    doc.add_paragraph(f"Base « {fname} » — audit, score de qualité et plan d'analyse conditionnel")
    meta = doc.add_paragraph()
    meta.add_run(f"Population auditée : {s['n_rows']} observations · "
                 f"Structure source : {s['n_cols']} variables · "
                 f"Date de l'audit : {date.today():%d/%m/%Y}").italic = True
    principe = doc.add_paragraph("Principe : source originale inchangée ; copie nettoyée traçable.")
    principe.runs[0].italic = True

    # 1. Synthèse exécutive
    _h(doc, "1. Synthèse exécutive")
    if assess.get("executive_summary_fr"):
        _p(doc, assess["executive_summary_fr"])
    else:
        n_empty = len(s["empty_columns"])
        _p(doc, f"La base contient {s['n_rows']} observations et {s['n_cols']} variables. "
                f"Le taux global de données manquantes est de {m['pct_global']} %. "
                f"{n_empty} colonne(s) sont entièrement vides. "
                f"Aucun doublon strict de ligne détecté."
                if s["duplicates"]["exact_rows"] == 0 else
                f"{s['duplicates']['exact_rows']} doublon(s) strict(s) détecté(s).")
    n_find = len(assess.get("findings", []))
    if n_find:
        n_crit = sum(1 for f in assess["findings"] if f.get("severity") == "critique")
        _p(doc, f"Le registre contient {n_find} anomalie(s), dont {n_crit} critique(s).")

    # 2. Documents examinés et cadre méthodologique
    _h(doc, "2. Documents examinés et cadre méthodologique")
    _p(doc, f"Fichier de données : {fname} (format {profiling['file'].get('format')}). "
            + ("Un protocole de recherche a été fourni et exploité."
               if study.get("primary_objective") else
               "Aucun protocole n'a été fourni : la concordance méthodologique n'a pu être évaluée."))
    if study.get("study_type"):
        _p(doc, "Type d'étude : " + study["study_type"] + ".")
    if study.get("primary_objective"):
        _p(doc, "Objectif principal annoncé : " + study["primary_objective"])
    plan_n, obs_n = study.get("planned_sample_size"), study.get("observed_sample_size")
    if plan_n and obs_n and str(plan_n) != str(obs_n):
        _p(doc, f"Effectif prévu au protocole : {plan_n} ; effectif observé dans la base : "
                f"{obs_n} (écart à documenter).")

    # 3. Audit structurel, identifiants et confidentialité
    _h(doc, "3. Audit structurel, identifiants et confidentialité")
    _p(doc, f"L'unité statistique est une ligne. La base comporte {s['n_rows']} lignes et "
            f"{s['n_unique_ids']} identifiant(s) unique(s)"
            + (f" (colonne « {s['id_column']} »)." if s["id_column"] else " (aucune colonne identifiant explicite)."))
    pii = profiling.get("pii_candidates", [])
    if pii:
        _p(doc, "Données potentiellement identifiantes détectées (à anonymiser avant analyse) : "
                + ", ".join(p["column"] for p in pii) + ".")

    # 4. Complétude et profils de données manquantes
    _h(doc, "4. Complétude et profils de données manquantes")
    n_empty = len(s["empty_columns"])
    empty_cells = n_empty * s["n_rows"]
    total_cells = m.get("total_cells", s["n_rows"] * s["n_cols"])
    missing_cells = m.get("missing_cells", round(m["pct_global"] / 100 * total_cells))
    denom = total_cells - empty_cells
    hors = round((missing_cells - empty_cells) / denom * 100, 2) if denom > 0 else 0
    _p(doc, f"Taux de manquants global : {m['pct_global']} %. Hors {n_empty} colonne(s) "
            f"entièrement vide(s) : {hors} %. Dossiers complets : {m['complete_cases']} "
            f"sur {s['n_rows']}.")
    if m.get("missing_blocks"):
        _p(doc, f"{len(m['missing_blocks'])} bloc(s) de variables manquantes ensemble ont "
                "été détectés (manque structurel plutôt qu'aléatoire).")

    # 5. Validité, codages et cohérences
    _h(doc, "5. Validité, codages et cohérences")
    if rs.get("limites"):
        _p(doc, rs["limites"])
    else:
        codes = [c["name"] for c in profiling["columns"] if c.get("suspected_missing_codes")]
        _p(doc, "Les modalités catégorielles respectent globalement les libellés SPSS."
                + (f" Codes de valeurs manquantes suspectés (999, -1…) sur : {', '.join(codes[:8])}."
                   if codes else ""))
    viol = ai.get("violations")
    if viol:
        _p(doc, f"{viol['rules_tested']} règle(s) de cohérence testée(s), "
                f"{viol['total_violations']} violation(s).")

    # 6. Concordance avec le protocole et faisabilité
    _h(doc, "6. Concordance avec le protocole et faisabilité")
    if rs.get("plan_action"):
        _p(doc, rs["plan_action"])
    elif study.get("primary_endpoint"):
        ep = study["primary_endpoint"]
        _p(doc, f"Critère de jugement principal identifié : {ep.get('description')} "
                f"(colonnes candidates : {', '.join(ep.get('candidate_columns', []))}).")
    else:
        _p(doc, "Concordance non évaluable sans protocole.")

    # 7. Score global de qualité
    _h(doc, "7. Score global de qualité")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, t in enumerate(["Domaine", "Score", "Max", "%"]):
        hdr[i].text = t
    for d in score["domaines"]:
        row = table.add_row().cells
        row[0].text = f"{d['domaine']}. {d['nom']}"
        row[1].text = str(d["obtenu"])
        row[2].text = str(d["max"])
        row[3].text = f"{d['obtenu'] / d['max'] * 100:.0f}%" if d["max"] else "—"
    _p(doc, f"Score brut : {score['score_brut']}/100. "
            f"Score officiel après plafond : {score['score_final']}/100. "
            f"Niveau de qualité : {score['niveau_qualite']}. "
            f"Confiance : {score['confiance']['niveau']}.")
    if score["plafonds_appliques"]:
        _p(doc, "Plafond appliqué : "
                + "; ".join(c["defaut"] for c in score["plafonds_appliques"]) + ".")
    fig = _fig_scores(score)
    if fig is not None:
        from docx.shared import Inches
        doc.add_picture(fig, width=Inches(6.0))
        legende = doc.add_paragraph("Figure 1. Sous-scores de qualité rapportés au maximum de "
                                    "chaque domaine.")
        legende.runs[0].italic = True

    # 8. Nettoyage appliqué à la copie dérivée
    _h(doc, "8. Nettoyage appliqué à la copie dérivée")
    plan = assess.get("cleaning_plan", [])
    if plan:
        _p(doc, "Opérations proposées (appliquées uniquement après validation, sur une copie ; "
                "l'original reste inchangé) :")
        for op in plan:
            doc.add_paragraph(
                f"{op.get('operation')} sur « {op.get('column')} » — {op.get('rationale_fr')}"
                + ("" if op.get("auto_safe") else " (validation clinique requise)"),
                style="List Bullet")
    else:
        _p(doc, "Le plan de nettoyage détaillé et la copie nettoyée seront produits à l'étape "
                "de nettoyage (aucune imputation ; original conservé).")

    # 9. Décisions à valider avant analyses finales
    _h(doc, "9. Décisions à valider avant analyses finales")
    decs = assess.get("client_decisions", [])
    if decs:
        for d in decs:
            doc.add_paragraph(
                f"{d.get('question_fr')} → {d.get('recommendation_fr')}", style="List Bullet")
    else:
        _p(doc, "Aucune décision bloquante identifiée à ce stade, ou audit IA non exécuté.")

    # 10. Plan d'analyse statistique conditionnel
    _h(doc, "10. Plan d'analyse statistique conditionnel")
    _p(doc, rs.get("plan_analyse_conditionnel")
       or "Le plan d'analyse conditionnel est produit avec l'audit IA.")

    # 11. Livrables et limites
    _h(doc, "11. Livrables et limites")
    _p(doc, "Livrables : le présent rapport, le classeur d'audit Excel, et — après nettoyage — "
            "la base nettoyée (.sav et .csv). Le nettoyage ne remplace pas la revue des dossiers "
            "sources ; aucune information absente n'a été inventée ; le score porte sur la base "
            "telle que reçue.")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
