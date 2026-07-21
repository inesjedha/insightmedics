"""Test du rapport Word (M6) — construction déterministe, sans API."""

import io

from docx import Document
from tests.test_report_xlsx import SYNTH_AI, _profiling_min
from tests.test_score_engine import CLEAN, FULL_SI

from app.pipeline.report_docx import build_report
from app.pipeline.score_engine import compute_score


def test_rapport_word_11_sections():
    prof = _profiling_min()
    score = compute_score(CLEAN, FULL_SI)
    data = build_report(prof, score, SYNTH_AI)
    doc = Document(io.BytesIO(data))
    headings = [p.text for p in doc.paragraphs if p.style.name == "Heading 1"]
    # 11 sections numérotées
    assert len(headings) == 11
    assert headings[0].startswith("1. Synthèse")
    assert headings[-1].startswith("11. Livrables")
    # Le texte contient le score final et le nombre de patients
    full = "\n".join(p.text for p in doc.paragraphs)
    assert str(score["score_final"]) in full
    assert "observations" in full


def test_rapport_word_sans_ia_reste_coherent():
    prof = _profiling_min()
    score = compute_score(CLEAN)
    data = build_report(prof, score, None)
    doc = Document(io.BytesIO(data))
    headings = [p.text for p in doc.paragraphs if p.style.name == "Heading 1"]
    assert len(headings) == 11  # toutes les sections présentes même sans IA
